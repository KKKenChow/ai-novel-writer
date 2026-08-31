"""小说导出：Markdown / docx / EPUB（EPUB 为零依赖手工打包）"""
import os
import zipfile
from typing import Dict
from xml.sax.saxutils import escape


def _chapters_sorted(chapters: Dict) -> list:
    def _k(item):
        try:
            return int(item[0])
        except (ValueError, TypeError):
            return 0
    return sorted(chapters.items(), key=_k)


def build_markdown(novel_name: str, chapters: Dict) -> str:
    parts = [f"# {novel_name}\n"]
    for num, data in _chapters_sorted(chapters):
        parts.append(f"\n## 第{num}章 {data.get('title', '')}\n\n{data.get('content', '')}\n")
    return "".join(parts)


def build_docx(novel_name: str, chapters: Dict, path: str):
    """导出 docx（需要 python-docx）"""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise RuntimeError("导出 docx 需要安装 python-docx：pip install python-docx")
    doc = Document()
    doc.add_heading(novel_name, level=0)
    for num, data in _chapters_sorted(chapters):
        doc.add_heading(f"第{num}章 {data.get('title', '')}", level=1)
        for para in data.get("content", "").split("\n"):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
        doc.add_page_break()
    doc.save(path)
    return path


def _xhtml_page(title: str, body_paragraphs: list) -> str:
    paras = "".join(f"<p>{escape(p)}</p>" for p in body_paragraphs if p.strip())
    return f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml"><head><title>{escape(title)}</title></head>
<body><h1>{escape(title)}</h1>{paras}</body></html>"""


def build_epub(novel_name: str, chapters: Dict, path: str, author: str = "AI 创作"):
    """导出 EPUB（零依赖手工打包，含目录导航）"""
    items = _chapters_sorted(chapters)
    manifest_items = []
    spine_items = []
    nav_points = []
    files = {}

    for i, (num, data) in enumerate(items):
        fname = f"chap_{num}.xhtml"
        title = f"第{num}章 {data.get('title', '')}"
        files[f"OEBPS/{fname}"] = _xhtml_page(title, data.get("content", "").split("\n"))
        manifest_items.append(f'<item id="c{i}" href="{fname}" media-type="application/xhtml+xml"/>')
        spine_items.append(f'<itemref idref="c{i}"/>')
        nav_points.append(
            f'<navPoint id="nav{i}" playOrder="{i+1}">'
            f'<navLabel><text>{escape(title)}</text></navLabel>'
            f'<content src="{fname}"/></navPoint>')

    files["mimetype"] = "application/epub+zip"
    files["META-INF/container.xml"] = """<?xml version="1.0" encoding="utf-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
<rootfiles><rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/></rootfiles>
</container>"""
    files["OEBPS/content.opf"] = f"""<?xml version="1.0" encoding="utf-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="bid" version="2.0">
<metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
<dc:title>{escape(novel_name)}</dc:title><dc:creator>{escape(author)}</dc:creator>
<dc:language>zh-CN</dc:language><dc:identifier id="bid">ai-novel-{escape(novel_name)}</dc:identifier>
</metadata>
<manifest><item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>{''.join(manifest_items)}</manifest>
<spine toc="ncx">{''.join(spine_items)}</spine>
</package>"""
    files["OEBPS/toc.ncx"] = f"""<?xml version="1.0" encoding="utf-8"?>
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
<head><meta name="dtb:uid" content="ai-novel"/></head>
<docTitle><text>{escape(novel_name)}</text></docTitle>
<navMap>{''.join(nav_points)}</navMap>
</ncx>"""

    with zipfile.ZipFile(path, "w") as zf:
        # mimetype 必须第一个且不压缩（EPUB 规范）
        zf.writestr(zipfile.ZipInfo("mimetype"), files.pop("mimetype"), compress_type=zipfile.ZIP_STORED)
        for name, content in files.items():
            zf.writestr(name, content, compress_type=zipfile.ZIP_DEFLATED)
    return path
